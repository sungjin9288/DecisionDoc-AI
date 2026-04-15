"""docx_service — build an in-memory DOCX from rendered markdown docs.

No disk I/O is performed; ``build_docx`` returns raw bytes.

Government format (행안부 공문서 표준):
- A4 (210×297mm), 맑은 고딕 10.5pt, 줄간격 160%
- 상 30mm / 하 15mm / 좌우 20mm 여백
- 헤더(기관명·문서번호) / 푸터(페이지 번호)
- 공문서 헤더 블록: 문서번호 / 수신 / 경유 / 제목 / 붙임
- 결재란 표
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from app.services.export_labels import humanize_doc_type
from app.services.export_outline import summarize_export_docs
from app.services.markdown_utils import parse_markdown_blocks


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _set_run_font(run: Any, font_name: str, font_size_pt: float | None = None) -> None:
    """Set both Latin and East-Asian (CJK) font on a run."""
    run.font.name = font_name
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    # East-Asian font for Korean characters
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _set_default_font(doc: Document, font_name: str, font_size_pt: float) -> None:
    """Apply Korean-compatible default font to the Normal paragraph style."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    # East-Asian font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _set_line_spacing(para: Any, spacing_pct: int) -> None:
    """Set paragraph line spacing as a multiple (percentage / 100)."""
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = spacing_pct / 100.0


def _add_field_code(run: Any, field_code: str) -> None:
    """Insert a Word field code (PAGE, NUMPAGES) into a run."""
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" {field_code} "
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _add_page_number_footer(section: Any, font_name: str = "맑은 고딕") -> None:
    """Add centered 'n / total' page numbers to the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    run = para.add_run()
    _set_run_font(run, font_name)
    _add_field_code(run, "PAGE")

    run_sep = para.add_run(" / ")
    _set_run_font(run_sep, font_name)

    run2 = para.add_run()
    _set_run_font(run2, font_name)
    _add_field_code(run2, "NUMPAGES")


def _add_gov_header(section: Any, opts: Any) -> None:
    """Add a right-aligned header showing org name / classification."""
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    para.clear()
    parts = []
    if opts and opts.org_name:
        parts.append(opts.org_name)
    if opts and opts.classification:
        parts.append(opts.classification)
    if parts:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run("  ".join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        _set_run_font(run, opts.font_name if opts else "맑은 고딕")


def _add_bold_inline(paragraph: Any, text: str,
                     font_name: str = "맑은 고딕",
                     font_size_pt: float = 10.5) -> None:
    """Parse **bold** spans and add them with appropriate runs."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        _set_run_font(run, font_name, font_size_pt)


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_table_cell_text(
    cell: Any,
    text: str,
    *,
    font_name: str,
    font_size_pt: float,
    line_spacing_pct: int,
    bold: bool = False,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    _add_bold_inline(paragraph, text, font_name, font_size_pt)
    if bold:
        for run in paragraph.runs:
            run.bold = True
    _set_line_spacing(paragraph, line_spacing_pct)


def _add_markdown_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    font_name: str,
    font_size_pt: float,
    line_spacing_pct: int,
) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        _add_table_cell_text(
            cell,
            header,
            font_name=font_name,
            font_size_pt=font_size_pt,
            line_spacing_pct=line_spacing_pct,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_shading(cell, "EAEFF7")

    for row_idx, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            _add_table_cell_text(
                table.cell(row_idx, col),
                value,
                font_name=font_name,
                font_size_pt=font_size_pt,
                line_spacing_pct=line_spacing_pct,
            )


def _add_markdown_content(doc: Document, markdown: str,
                           font_name: str = "맑은 고딕",
                           font_size_pt: float = 10.5,
                           line_spacing_pct: int = 160) -> None:
    """Convert markdown lines to Word paragraphs/headings/lists."""
    for block in parse_markdown_blocks(markdown):
        block_type = block["type"]
        if block_type == "heading":
            level = min(int(block.get("level", 1)), 3)
            p = doc.add_heading(block["text"], level=level)
        elif block_type == "list_item":
            p = doc.add_paragraph(style="List Bullet")
            _add_bold_inline(p, block["text"], font_name, font_size_pt)
            _set_line_spacing(p, line_spacing_pct)
            continue
        elif block_type == "table":
            _add_markdown_table(
                doc,
                block["headers"],
                block["rows"],
                font_name=font_name,
                font_size_pt=font_size_pt,
                line_spacing_pct=line_spacing_pct,
            )
            continue
        elif block_type == "hr":
            p = doc.add_paragraph()
        elif block_type == "blank":
            p = doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _add_bold_inline(p, block["text"], font_name, font_size_pt)
            _set_line_spacing(p, line_spacing_pct)
            continue
        _set_line_spacing(p, line_spacing_pct)


def _add_export_cover_page(
    doc: Document,
    *,
    title: str,
    docs: list[dict[str, Any]],
    font_name: str,
    font_size_pt: float,
    line_spacing_pct: int,
) -> None:
    summaries = summarize_export_docs(docs)
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    run.bold = True
    _set_run_font(run, font_name, font_size_pt + 9)
    _set_line_spacing(cover, 130)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("완성형 문서 패키지")
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x5F, 0x63, 0x77)
    _set_run_font(subtitle_run, font_name, font_size_pt + 1)
    _set_line_spacing(subtitle, line_spacing_pct)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary.add_run(f"총 {len(docs)}개 문서를 하나의 제출 패키지로 정리했습니다.")
    _set_run_font(summary_run, font_name, font_size_pt)
    _set_line_spacing(summary, line_spacing_pct)

    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("문서 구성")
    heading_run.bold = True
    _set_run_font(heading_run, font_name, font_size_pt + 1.5)

    for idx, item in enumerate(docs, start=1):
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_bold_inline(
            line,
            f"**{idx}.** {humanize_doc_type(item.get('doc_type', 'document'))}",
            font_name,
            font_size_pt,
        )
        _set_line_spacing(line, line_spacing_pct)

    doc.add_paragraph()
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_heading.add_run("핵심 검토 포인트")
    summary_run.bold = True
    _set_run_font(summary_run, font_name, font_size_pt + 1.5)

    table = doc.add_table(rows=len(summaries) + 1, cols=3)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    header_cells = ["문서", "핵심 메시지", "구성 특징"]
    for col, header in enumerate(header_cells):
        cell = table.cell(0, col)
        _add_table_cell_text(
            cell,
            header,
            font_name=font_name,
            font_size_pt=font_size_pt,
            line_spacing_pct=line_spacing_pct,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_shading(cell, "EAEFF7")

    for row_idx, summary in enumerate(summaries, start=1):
        _add_table_cell_text(
            table.cell(row_idx, 0),
            summary["label"],
            font_name=font_name,
            font_size_pt=font_size_pt,
            line_spacing_pct=line_spacing_pct,
            bold=True,
        )
        _add_table_cell_text(
            table.cell(row_idx, 1),
            summary["lead"],
            font_name=font_name,
            font_size_pt=font_size_pt,
            line_spacing_pct=line_spacing_pct,
        )
        _add_table_cell_text(
            table.cell(row_idx, 2),
            f"{summary['sections']} / {summary['metrics']}",
            font_name=font_name,
            font_size_pt=font_size_pt,
            line_spacing_pct=line_spacing_pct,
        )

    doc.add_page_break()


def _add_doc_section_intro(
    doc: Document,
    *,
    doc_type: str,
    index: int,
    total: int,
    lead: str,
    section_hint: str,
    metrics: str,
    font_name: str,
    font_size_pt: float,
    line_spacing_pct: int,
) -> None:
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.LEFT
    badge_run = badge.add_run(f"문서 {index:02d} / {total:02d}")
    badge_run.bold = True
    badge_run.font.color.rgb = RGBColor(0x5B, 0x63, 0xD3)
    _set_run_font(badge_run, font_name, font_size_pt)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(humanize_doc_type(doc_type))
    title_run.bold = True
    _set_run_font(title_run, font_name, font_size_pt + 5)
    _set_line_spacing(title_para, 130)

    desc_para = doc.add_paragraph()
    _add_bold_inline(
        desc_para,
        lead,
        font_name,
        font_size_pt,
    )
    for run in desc_para.runs:
        run.font.color.rgb = RGBColor(0x5F, 0x63, 0x77)
    _set_line_spacing(desc_para, line_spacing_pct)

    meta_para = doc.add_paragraph()
    _add_bold_inline(
        meta_para,
        f"핵심 섹션: {section_hint} / 구성 특징: {metrics}",
        font_name,
        font_size_pt - 0.3,
    )
    for run in meta_para.runs:
        run.font.color.rgb = RGBColor(0x6E, 0x75, 0x91)
    _set_line_spacing(meta_para, line_spacing_pct)

    separator = doc.add_paragraph("─" * 48)
    separator.runs[0].font.color.rgb = RGBColor(0xB9, 0xC2, 0xD6)
    _set_run_font(separator.runs[0], font_name, font_size_pt - 1)
    _set_line_spacing(separator, 100)


# ---------------------------------------------------------------------------
# Government format helpers
# ---------------------------------------------------------------------------

def _add_gov_doc_header_block(doc: Document, title: str, opts: Any) -> None:
    """행안부 공문서 헤더 블록: 기관명 / 문서번호 / 수신 / 경유 / 제목 / 붙임."""
    fn = opts.font_name
    fs = opts.font_size_pt

    # 기관명 (가운데 정렬, 굵게)
    if opts.org_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(opts.org_name)
        run.bold = True
        _set_run_font(run, fn, fs + 2)
        _set_line_spacing(p, opts.line_spacing_pct)

    # 부서명 (가운데 정렬)
    if opts.dept_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(opts.dept_name)
        _set_run_font(run, fn, fs)
        _set_line_spacing(p, opts.line_spacing_pct)

    if opts.org_name or opts.dept_name:
        doc.add_paragraph()

    # 구분선
    p = doc.add_paragraph("─" * 60)
    p.runs[0].font.size = Pt(8)
    _set_line_spacing(p, 100)

    # 문서번호
    if opts.doc_number:
        p = doc.add_paragraph()
        r1 = p.add_run("문서번호: ")
        r1.bold = True
        _set_run_font(r1, fn, fs)
        r2 = p.add_run(opts.doc_number)
        _set_run_font(r2, fn, fs)
        _set_line_spacing(p, opts.line_spacing_pct)

    # 수신
    if opts.recipient:
        p = doc.add_paragraph()
        r1 = p.add_run("수\u2003\u2003신: ")
        r1.bold = True
        _set_run_font(r1, fn, fs)
        r2 = p.add_run(opts.recipient)
        _set_run_font(r2, fn, fs)
        _set_line_spacing(p, opts.line_spacing_pct)

    # 경유
    if opts.via:
        p = doc.add_paragraph()
        r1 = p.add_run("경\u2003\u2003유: ")
        r1.bold = True
        _set_run_font(r1, fn, fs)
        r2 = p.add_run(opts.via)
        _set_run_font(r2, fn, fs)
        _set_line_spacing(p, opts.line_spacing_pct)

    # 제목 (굵게)
    p = doc.add_paragraph()
    r1 = p.add_run("제\u2003\u2003목: ")
    r1.bold = True
    _set_run_font(r1, fn, fs)
    r2 = p.add_run(title)
    r2.bold = True
    _set_run_font(r2, fn, fs)
    _set_line_spacing(p, opts.line_spacing_pct)

    doc.add_paragraph()

    # 붙임
    if opts.attachments:
        p = doc.add_paragraph()
        r1 = p.add_run("붙\u2003\u2003임: ")
        r1.bold = True
        _set_run_font(r1, fn, fs)
        for i, att in enumerate(opts.attachments, 1):
            r2 = p.add_run(f"{i}. {att}  ")
            _set_run_font(r2, fn, fs)
        _set_line_spacing(p, opts.line_spacing_pct)

    doc.add_paragraph()


def _add_approval_block(doc: Document, opts: Any) -> None:
    """행안부 공문서 결재란 표를 문서 말미에 추가."""
    approvers: list[tuple[str, str]] = []
    if opts.drafter:
        approvers.append(("기  안", opts.drafter))
    if opts.reviewer:
        approvers.append(("검  토", opts.reviewer))
    if opts.approver:
        approvers.append(("결  재", opts.approver))

    if not approvers:
        return

    doc.add_paragraph()
    p_label = doc.add_paragraph()
    r = p_label.add_run("결  재")
    r.bold = True
    _set_run_font(r, opts.font_name, opts.font_size_pt)
    p_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = doc.add_table(rows=2, cols=len(approvers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass  # style may not exist in all templates

    # Row 0: role titles
    for col, (role, _) in enumerate(approvers):
        cell = table.cell(0, col)
        cell.text = role
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                _set_run_font(run, opts.font_name, opts.font_size_pt)

    # Row 1: names with signature space
    for col, (_, name) in enumerate(approvers):
        cell = table.cell(1, col)
        cell.text = ""
        # spacing for signature area
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(30)
        p.paragraph_format.space_after = Pt(5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        _set_run_font(run, opts.font_name, opts.font_size_pt)

    # Right-align the table
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    tbl_pr.append(jc)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_docx(
    docs: list[dict[str, Any]],
    title: str,
    gov_options: Any | None = None,
) -> bytes:
    """Build a DOCX from a list of rendered docs.

    Args:
        docs: List of {"doc_type": str, "markdown": str}.
        title: Document title.
        gov_options: Optional ``GovDocOptions`` dataclass instance.  When
            ``gov_options.is_government_format`` is ``True`` the full 행안부
            공문서 standard layout is applied (margins, header block, approval
            table).  Without gov_options the document uses sensible defaults
            that still produce an A4 document with Korean font.

    Returns:
        Raw bytes of the .docx file.
    """
    opts = gov_options

    # Resolve layout parameters (with fallbacks for no gov_options)
    top_mm    = opts.top_margin_mm    if opts else 30
    bot_mm    = opts.bottom_margin_mm if opts else 15
    left_mm   = opts.left_margin_mm   if opts else 20
    right_mm  = opts.right_margin_mm  if opts else 20
    font_name = opts.font_name        if opts else "맑은 고딕"
    font_size = opts.font_size_pt     if opts else 10.5
    spacing   = opts.line_spacing_pct if opts else 160

    doc = Document()

    # ── Page layout ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Mm(210)   # A4
    section.page_height = Mm(297)   # A4
    section.top_margin    = Mm(top_mm)
    section.bottom_margin = Mm(bot_mm)
    section.left_margin   = Mm(left_mm)
    section.right_margin  = Mm(right_mm)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    # ── Default font ──────────────────────────────────────────────────────
    _set_default_font(doc, font_name, font_size)

    # ── Header / Footer ───────────────────────────────────────────────────
    _add_gov_header(section, opts)
    _add_page_number_footer(section, font_name)

    # ── Document body ─────────────────────────────────────────────────────
    summaries = summarize_export_docs(docs)
    if opts and opts.is_government_format:
        _add_gov_doc_header_block(doc, title, opts)
    else:
        _add_export_cover_page(
            doc,
            title=title,
            docs=docs,
            font_name=font_name,
            font_size_pt=font_size,
            line_spacing_pct=spacing,
        )

    for i, d in enumerate(docs):
        if i > 0:
            doc.add_page_break()
        if not (opts and opts.is_government_format):
            summary = summaries[i]
            _add_doc_section_intro(
                doc,
                doc_type=str(d.get("doc_type", "document")),
                index=i + 1,
                total=len(docs),
                lead=summary["lead"],
                section_hint=summary["sections"],
                metrics=summary["metrics"],
                font_name=font_name,
                font_size_pt=font_size,
                line_spacing_pct=spacing,
            )
        _add_markdown_content(doc, d.get("markdown", ""), font_name, font_size, spacing)

    # ── Approval block (공문서 전용) ──────────────────────────────────────
    if opts and opts.is_government_format:
        _add_approval_block(doc, opts)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
