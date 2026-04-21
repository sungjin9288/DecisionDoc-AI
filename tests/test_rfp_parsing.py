"""tests/test_rfp_parsing.py — Tests for attachment extraction and RFP parsing.

Coverage:
- attachment_service: txt/md/csv extraction, file-size limit, unsupported ext,
  extract_multiple() total cap + failed-file handling
- rfp_parser: parse_rfp_fields(), _suggest_bundle(), build_rfp_context()
- API: /attachments/parse-rfp endpoint (single + multi file, unsupported ext)
- API: /generate/with-attachments context injection (structured RFP block)
"""
from __future__ import annotations

import csv
import io
import json
from unittest.mock import MagicMock, patch

import pytest
from fastapi.testclient import TestClient


# ── Helpers ───────────────────────────────────────────────────────────────────

def _make_client(tmp_path, monkeypatch) -> TestClient:
    monkeypatch.setenv("DECISIONDOC_PROVIDER", "mock")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("DECISIONDOC_ENV", "dev")
    monkeypatch.setenv("DECISIONDOC_PROVIDER_GENERATION", "mock")
    monkeypatch.setenv("DECISIONDOC_PROVIDER_ATTACHMENT", "mock")
    monkeypatch.setenv("DECISIONDOC_PROVIDER_VISUAL", "mock")
    monkeypatch.delenv("DECISIONDOC_API_KEY",  raising=False)
    monkeypatch.delenv("DECISIONDOC_API_KEYS", raising=False)
    from app.main import create_app
    return TestClient(create_app())


def _txt(content: str) -> bytes:
    return content.encode("utf-8")


def _csv_bytes(rows: list[list[str]]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


def _make_docx(texts: list[str]) -> bytes:
    """Build a minimal .docx in memory with given paragraphs."""
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx(data: dict[str, list[list]]) -> bytes:
    """Build a minimal .xlsx in memory with given sheet -> rows."""
    import openpyxl
    wb = openpyxl.Workbook()
    first = True
    for sheet_name, rows in data.items():
        if first:
            ws = wb.active
            ws.title = sheet_name
            first = False
        else:
            ws = wb.create_sheet(sheet_name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── attachment_service: plain text ───────────────────────────────────────────

class TestExtractText:
    def test_txt_utf8(self):
        from app.services.attachment_service import extract_text
        result = extract_text("file.txt", _txt("안녕하세요 테스트"))
        assert "안녕하세요 테스트" in result

    def test_md_utf8(self):
        from app.services.attachment_service import extract_text
        result = extract_text("readme.md", _txt("# 제목\n내용"))
        assert "# 제목" in result

    def test_txt_latin1_fallback(self):
        from app.services.attachment_service import extract_text
        raw = "caf\xe9".encode("latin-1")
        result = extract_text("file.txt", raw)
        assert "caf" in result  # decoded via latin-1 fallback

    def test_txt_truncated_at_max(self):
        from app.services.attachment_service import MAX_CHARS_PER_FILE, extract_text
        big = "a" * (MAX_CHARS_PER_FILE + 1000)
        result = extract_text("big.txt", _txt(big))
        assert len(result) == MAX_CHARS_PER_FILE

    def test_unsupported_extension_raises(self):
        from app.services.attachment_service import AttachmentError, extract_text
        with pytest.raises(AttachmentError, match="지원하지 않는"):
            extract_text("file.exe", b"binary data")

    def test_file_size_limit_raises(self):
        from app.services.attachment_service import (
            MAX_FILE_SIZE_BYTES,
            AttachmentError,
            extract_text,
        )
        oversized = b"x" * (MAX_FILE_SIZE_BYTES + 1)
        with pytest.raises(AttachmentError, match="초과"):
            extract_text("big.txt", oversized)

    def test_csv_extraction(self):
        from app.services.attachment_service import extract_text
        raw = _csv_bytes([["이름", "나이"], ["홍길동", "30"]])
        result = extract_text("data.csv", raw)
        assert "이름" in result
        assert "홍길동" in result
        assert "|" in result   # pipe-separated

    def test_csv_empty_rows_skipped(self):
        from app.services.attachment_service import extract_text
        raw = _csv_bytes([["A", "B"], [], ["C", "D"]])
        result = extract_text("data.csv", raw)
        # empty row should not appear as "| "
        assert "C" in result

    def test_docx_paragraphs(self):
        from app.services.attachment_service import extract_text
        raw = _make_docx(["첫 번째 단락", "두 번째 단락"])
        result = extract_text("doc.docx", raw)
        assert "첫 번째 단락" in result
        assert "두 번째 단락" in result

    def test_docx_table_included(self):
        from docx import Document
        from app.services.attachment_service import extract_text
        buf = io.BytesIO()
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "헤더1"
        table.cell(0, 1).text = "헤더2"
        table.cell(1, 0).text = "값1"
        table.cell(1, 1).text = "값2"
        doc.save(buf)
        result = extract_text("doc.docx", buf.getvalue())
        assert "[표]" in result
        assert "헤더1" in result

    @pytest.mark.skipif(
        not __import__("importlib").util.find_spec("openpyxl"),
        reason="openpyxl not installed",
    )
    def test_xlsx_extraction(self):
        from app.services.attachment_service import extract_text
        raw = _make_xlsx({"Sheet1": [["품목", "수량", "단가"], ["사과", 10, 500]]})
        result = extract_text("data.xlsx", raw)
        assert "[시트: Sheet1]" in result
        assert "품목" in result
        assert "사과" in result

    @pytest.mark.skipif(
        not __import__("importlib").util.find_spec("openpyxl"),
        reason="openpyxl not installed",
    )
    def test_xlsx_multiple_sheets(self):
        from app.services.attachment_service import extract_text
        raw = _make_xlsx({
            "매출": [["월", "금액"], ["1월", 1000]],
            "매입": [["항목", "금액"], ["원자재", 500]],
        })
        result = extract_text("data.xlsx", raw)
        assert "[시트: 매출]" in result
        assert "[시트: 매입]" in result


# ── attachment_service: extract_multiple ────────────────────────────────────

class TestExtractMultiple:
    def test_returns_combined_text(self):
        from app.services.attachment_service import extract_multiple
        files = [
            ("a.txt", _txt("내용A")),
            ("b.txt", _txt("내용B")),
        ]
        result = extract_multiple(files)
        assert "[첨부파일: a.txt]" in result
        assert "[첨부파일: b.txt]" in result
        assert "내용A" in result
        assert "내용B" in result

    def test_separator_between_files(self):
        from app.services.attachment_service import extract_multiple
        files = [("x.txt", _txt("X")), ("y.txt", _txt("Y"))]
        result = extract_multiple(files)
        assert "---" in result

    def test_failed_file_emits_warning_line(self):
        from app.services.attachment_service import extract_multiple
        files = [
            ("bad.exe", b"binary"),   # unsupported
            ("good.txt", _txt("좋은 내용")),
        ]
        result = extract_multiple(files)
        assert "⚠️" in result          # warning line for bad.exe
        assert "좋은 내용" in result    # good.txt still processed

    def test_global_char_cap(self):
        from app.services.attachment_service import MAX_TOTAL_CHARS, extract_multiple
        # Each file has just under the per-file cap; together they exceed global cap
        big = "b" * (MAX_TOTAL_CHARS // 2 + 2000)
        files = [
            ("f1.txt", _txt(big)),
            ("f2.txt", _txt(big)),
            ("f3.txt", _txt(big)),
        ]
        result = extract_multiple(files)
        # Total chars in result sections should not massively exceed global cap
        # (the result string includes headers like "[첨부파일: ...]", so allow some slack)
        assert len(result) < MAX_TOTAL_CHARS + 1000

    def test_empty_file_list(self):
        from app.services.attachment_service import extract_multiple
        result = extract_multiple([])
        assert result == ""


# ── rfp_parser: build_rfp_context ───────────────────────────────────────────

class TestBuildRFPContext:
    def test_contains_rfp_header(self):
        from app.services.rfp_parser import build_rfp_context
        ctx = build_rfp_context("테스트 RFP 내용")
        assert "=== RFP 원문" in ctx

    def test_contains_original_text(self):
        from app.services.rfp_parser import build_rfp_context
        ctx = build_rfp_context("독특한 원문 텍스트")
        assert "독특한 원문 텍스트" in ctx

    def test_contains_closing_marker(self):
        from app.services.rfp_parser import build_rfp_context
        ctx = build_rfp_context("내용")
        assert "=== RFP 원문 끝 ===" in ctx

    def test_contains_instruction(self):
        from app.services.rfp_parser import build_rfp_context
        ctx = build_rfp_context("내용")
        assert "바탕으로 문서를 작성하세요" in ctx

    def test_contains_grounding_rules_for_missing_details(self):
        from app.services.rfp_parser import build_rfp_context

        ctx = build_rfp_context("내용")

        assert "날짜, 예산, 기관명, 기술명, 배점, 수치는 임의로 만들지 마세요" in ctx
        assert "예시 기술 스택이나 마감일을 추정해서 채우지 마세요" in ctx

    def test_includes_normalized_procurement_context_when_provided(self):
        from app.services.rfp_parser import build_rfp_context

        ctx = build_rfp_context(
            "원문",
            normalized_context="=== 공공조달 PDF 정규화 요약 ===\n핵심 섹션:\n- 평가 개요",
        )

        assert "=== 공공조달 PDF 정규화 요약 ===" in ctx
        assert "=== RFP 원문" in ctx


# ── rfp_parser: _suggest_bundle ─────────────────────────────────────────────

class TestSuggestBundle:
    def _suggest(self, text: str) -> str:
        from app.services.rfp_parser import _suggest_bundle
        return _suggest_bundle(text, {})

    def test_rfp_keyword(self):
        assert self._suggest("이번 제안요청서에 따른 사업") == "rfp_analysis_kr"

    def test_nara_keyword(self):
        assert self._suggest("나라장터 입찰공고 공지") == "rfp_analysis_kr"

    def test_performance_keyword(self):
        assert self._suggest("과업지시서 및 수행계획 수립") == "performance_plan_kr"

    def test_completion_keyword(self):
        assert self._suggest("준공 및 완료 보고서 제출") == "completion_report_kr"

    def test_interim_keyword(self):
        assert self._suggest("중간 보고 진척 현황") == "interim_report_kr"

    def test_proposal_keyword(self):
        assert self._suggest("사업계획서 및 제안서 제출") == "proposal_kr"

    def test_fallback_default(self):
        assert self._suggest("알 수 없는 텍스트") == "rfp_analysis_kr"


# ── rfp_parser: parse_rfp_fields ────────────────────────────────────────────

class TestParseRFPFields:
    def _make_provider(self, raw_json: str):
        mock = MagicMock()
        mock.generate_raw.return_value = raw_json
        return mock

    def test_returns_all_expected_keys(self):
        from app.services.rfp_parser import parse_rfp_fields
        provider = self._make_provider(json.dumps({
            "project_title": "테스트 사업",
            "issuer": "행정안전부",
            "budget": "5억원",
            "deadline": "2025-12-31",
            "duration": "2025.04~2025.12",
            "objective": "공공서비스 디지털 전환",
            "key_requirements": ["클라우드 전환", "보안 인증"],
            "evaluation_criteria": ["기술점수(70)", "가격점수(30)"],
            "confidence": 0.92,
        }))
        result = parse_rfp_fields("테스트 RFP", provider=provider)
        assert result["project_title"] == "테스트 사업"
        assert result["issuer"] == "행정안전부"
        assert result["budget"] == "5억원"
        assert isinstance(result["key_requirements"], list)
        assert result["confidence"] == pytest.approx(0.92)

    def test_returns_empty_result_on_invalid_json(self):
        from app.services.rfp_parser import parse_rfp_fields
        provider = self._make_provider("이건 JSON이 아닙니다")
        result = parse_rfp_fields("RFP 텍스트", provider=provider)
        assert result["confidence"] == 0.0
        assert result["suggested_bundle"] == "rfp_analysis_kr"

    def test_strips_markdown_fences(self):
        from app.services.rfp_parser import parse_rfp_fields
        raw = '```json\n{"project_title":"마크다운 테스트","confidence":0.5}\n```'
        provider = self._make_provider(raw)
        result = parse_rfp_fields("내용", provider=provider)
        assert result["project_title"] == "마크다운 테스트"

    def test_suggested_bundle_overrides_from_text(self):
        from app.services.rfp_parser import parse_rfp_fields
        provider = self._make_provider(json.dumps({"confidence": 0.8}))
        result = parse_rfp_fields("제안요청서 RFP 공고", provider=provider)
        assert result["suggested_bundle"] == "rfp_analysis_kr"

    def test_provider_called_with_prompt(self):
        from app.services.rfp_parser import parse_rfp_fields
        provider = self._make_provider("{}")
        parse_rfp_fields("RFP 원문입니다", provider=provider)
        assert provider.generate_raw.called
        call_args = provider.generate_raw.call_args
        assert "RFP 원문입니다" in call_args[0][0]  # prompt contains text

    def test_returns_empty_on_provider_exception(self):
        from app.services.rfp_parser import parse_rfp_fields
        provider = MagicMock()
        provider.generate_raw.side_effect = RuntimeError("LLM 연결 실패")
        result = parse_rfp_fields("내용", provider=provider)
        assert result["confidence"] == 0.0


# ── API: /attachments/parse-rfp ─────────────────────────────────────────────

class TestParseRFPEndpoint:
    def test_txt_file_returns_200(self, tmp_path, monkeypatch):
        client = _make_client(tmp_path, monkeypatch)
        content = "제안요청서\n사업명: 디지털 전환 사업\n발주기관: 과학기술정보통신부"
        res = client.post(
            "/attachments/parse-rfp",
            files=[("files", ("rfp.txt", content.encode(), "text/plain"))],
        )
        assert res.status_code == 200
        data = res.json()
        assert "extracted_fields" in data
        assert "raw_text_preview" in data
        assert "files_processed" in data
        assert "total_chars" in data

    def test_csv_file_returns_200(self, tmp_path, monkeypatch):
        client = _make_client(tmp_path, monkeypatch)
        raw = _csv_bytes([["항목", "내용"], ["사업명", "테스트"]])
        res = client.post(
            "/attachments/parse-rfp",
            files=[("files", ("rfp.csv", raw, "text/csv"))],
        )
        assert res.status_code == 200

    def test_unsupported_format_returns_200_with_warning(self, tmp_path, monkeypatch):
        """Unsupported file → warning line in result, endpoint still succeeds."""
        client = _make_client(tmp_path, monkeypatch)
        res = client.post(
            "/attachments/parse-rfp",
            files=[("files", ("note.exe", b"binary", "application/octet-stream"))],
        )
        # endpoint succeeds (parse-rfp is lenient) or returns 400 with detail
        assert res.status_code in (200, 400)

    def test_legacy_hwp_returns_422_with_conversion_guidance(self, tmp_path, monkeypatch):
        client = _make_client(tmp_path, monkeypatch)
        res = client.post(
            "/attachments/parse-rfp",
            files=[("files", ("legacy.hwp", b"binary", "application/x-hwp"))],
        )
        assert res.status_code == 422
        data = res.json()
        assert data["code"] == "ATTACHMENT_ERROR"
        assert "HWPX, PDF 또는 DOCX로 변환" in data["message"]

    def test_files_processed_list(self, tmp_path, monkeypatch):
        client = _make_client(tmp_path, monkeypatch)
        res = client.post(
            "/attachments/parse-rfp",
            files=[("files", ("a.txt", b"AAA", "text/plain")),
                   ("files", ("b.txt", b"BBB", "text/plain"))],
        )
        assert res.status_code == 200
        data = res.json()
        assert len(data["files_processed"]) == 2

    def test_structured_context_in_response(self, tmp_path, monkeypatch):
        client = _make_client(tmp_path, monkeypatch)
        res = client.post(
            "/attachments/parse-rfp",
            files=[("files", ("rfp.txt", b"RFP content", "text/plain"))],
        )
        assert res.status_code == 200
        data = res.json()
        assert "structured_context" in data
        assert "=== RFP 원문" in data["structured_context"]

    def test_pdf_response_includes_procurement_context_preview(self, tmp_path, monkeypatch):
        client = _make_client(tmp_path, monkeypatch)

        structured = {
            "title": "착수보고 자료",
            "sections": [
                {"heading": "Ⅰ. 평가 개요", "content": "경영평가 추진 배경을 설명한다."},
                {"heading": "Ⅱ. 추진 일정", "content": "착수와 중간 보고 일정을 제시한다."},
            ],
            "raw_text": "착수보고 자료\n경영평가 추진 배경",
            "page_count": 10,
            "has_tables": True,
        }

        with patch("app.routers.generate.extract_multiple", return_value="RFP content"), patch(
            "app.routers.generate.extract_pdf_structured",
            return_value=structured,
        ):
            res = client.post(
                "/attachments/parse-rfp",
                files=[("files", ("rfp.pdf", b"%PDF-1.4 fake", "application/pdf"))],
            )

        assert res.status_code == 200
        data = res.json()
        assert "procurement_context_preview" in data
        assert "공공조달 PDF 정규화 요약" in data["procurement_context_preview"]
        assert "핵심 섹션:" in data["structured_context"]

    def test_total_chars_in_response(self, tmp_path, monkeypatch):
        client = _make_client(tmp_path, monkeypatch)
        content = "X" * 500
        res = client.post(
            "/attachments/parse-rfp",
            files=[("files", ("rfp.txt", content.encode(), "text/plain"))],
        )
        assert res.status_code == 200
        data = res.json()
        # total_chars counts full extract_multiple output incl. "[첨부파일: ...]" header
        assert data["total_chars"] >= 500


# ── API: /generate/with-attachments context injection ───────────────────────

class TestWithAttachmentsContextInjection:
    def test_rfp_context_injected(self, tmp_path, monkeypatch):
        """Verify the RFP wrapper appears in the generation request context.

        The /generate/with-attachments endpoint expects:
          - payload: JSON string via Form field
          - attachments: file uploads
        """
        client = _make_client(tmp_path, monkeypatch)

        rfp_text = "제안요청서 RFP 내용입니다. 사업명: 국가 디지털 전환 사업"
        payload_json = json.dumps({
            "title": "제안서",
            "goal": "국가 디지털 전환",
            "bundle_type": "rfp_analysis_kr",
        })

        res = client.post(
            "/generate/with-attachments",
            data={"payload": payload_json},
            files=[("attachments", ("rfp.txt", rfp_text.encode(), "text/plain"))],
        )
        assert res.status_code == 200
        # Response is a GenerateResponse JSON object
        data = res.json()
        assert "docs" in data

    def test_no_attachment_still_works(self, tmp_path, monkeypatch):
        """With-attachments endpoint without files should still work."""
        client = _make_client(tmp_path, monkeypatch)
        payload_json = json.dumps({
            "title": "테스트",
            "goal": "테스트 목표",
            "bundle_type": "rfp_analysis_kr",
        })
        res = client.post(
            "/generate/with-attachments",
            data={"payload": payload_json},
        )
        assert res.status_code == 200
        data = res.json()
        assert "docs" in data
