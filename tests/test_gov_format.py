"""tests/test_gov_format.py — Government document format compliance tests.

Covers:
  A. GovDocOptions dataclass defaults and field types
  B. DOCX: A4 page size, Korean font, margins, backward-compat
  C. DOCX: gov doc header block + approval block when is_government_format=True
  D. HWP: proper style definitions in header.xml
  E. HWP: page layout (secPr) with A4 size and configurable margins
  F. HWP: gov doc header paragraphs when is_government_format=True
  G. PDF: HTML structure contains margin CSS
  H. PDF: gov header block / approval block HTML when is_government_format=True
  I. API: EditedExportRequest accepts gov_options field
  J. API: /generate/export-edited passes gov_options to DOCX/HWP
  K. _resolve_gov_options helper
"""
from __future__ import annotations

import io
import json
import zipfile

import pytest
from docx import Document
from docx.shared import Mm

# --- import our modules ---
from app.schemas import GovDocOptions, EditedExportRequest, EditedDocInput
from app.services.docx_service import build_docx
from app.services.hwp_service import build_hwp, _styles_xml, _section_xml, _A4_W, _A4_H
from app.services.pdf_service import _build_css, _gov_header_block_html, _approval_block_html, _render_html


# ===========================================================================
# A. GovDocOptions dataclass
# ===========================================================================

class TestGovDocOptionsDefaults:
    def test_is_government_format_default_false(self):
        opts = GovDocOptions()
        assert opts.is_government_format is False

    def test_top_margin_default(self):
        assert GovDocOptions().top_margin_mm == 30

    def test_bottom_margin_default(self):
        assert GovDocOptions().bottom_margin_mm == 15

    def test_left_right_margin_default(self):
        opts = GovDocOptions()
        assert opts.left_margin_mm == 20
        assert opts.right_margin_mm == 20

    def test_font_defaults(self):
        opts = GovDocOptions()
        assert opts.font_name == "맑은 고딕"
        assert opts.font_size_pt == 10.5
        assert opts.line_spacing_pct == 160

    def test_attachments_default_empty_list(self):
        opts = GovDocOptions()
        assert opts.attachments == []

    def test_custom_values(self):
        opts = GovDocOptions(
            doc_number="행안부-0001",
            recipient="수신자 참조",
            org_name="행정안전부",
            is_government_format=True,
        )
        assert opts.doc_number == "행안부-0001"
        assert opts.recipient == "수신자 참조"
        assert opts.org_name == "행정안전부"
        assert opts.is_government_format is True


# ===========================================================================
# B. DOCX — A4 page size, Korean font, margins, backward-compat
# ===========================================================================

class TestDocxPageLayout:
    _DOCS = [{"doc_type": "test", "markdown": "# 제목\n\n본문 내용입니다."}]

    def _load(self, gov_options=None):
        raw = build_docx(self._DOCS, title="테스트 문서", gov_options=gov_options)
        return Document(io.BytesIO(raw))

    def test_backward_compat_no_gov_options(self):
        """build_docx still works when gov_options is None."""
        raw = build_docx(self._DOCS, title="레거시 문서")
        assert len(raw) > 100  # file was produced

    def test_a4_width(self):
        doc = self._load()
        section = doc.sections[0]
        assert abs(section.page_width.mm - 210) < 1

    def test_a4_height(self):
        doc = self._load()
        section = doc.sections[0]
        assert abs(section.page_height.mm - 297) < 1

    def test_default_top_margin_30mm(self):
        doc = self._load()
        assert abs(doc.sections[0].top_margin.mm - 30) < 1

    def test_default_bottom_margin_15mm(self):
        doc = self._load()
        assert abs(doc.sections[0].bottom_margin.mm - 15) < 1

    def test_custom_margins(self):
        opts = GovDocOptions(top_margin_mm=40, bottom_margin_mm=25,
                              left_margin_mm=25, right_margin_mm=25)
        doc = self._load(gov_options=opts)
        s = doc.sections[0]
        assert abs(s.top_margin.mm - 40) < 1
        assert abs(s.bottom_margin.mm - 25) < 1
        assert abs(s.left_margin.mm - 25) < 1

    def test_returns_bytes(self):
        raw = build_docx(self._DOCS, title="X")
        assert isinstance(raw, bytes)


# ===========================================================================
# C. DOCX — gov header block + approval block
# ===========================================================================

class TestDocxGovFormat:
    _DOCS = [{"doc_type": "t", "markdown": "## 배경\n\n내용"}]

    def _full_text(self, gov_options=None):
        raw = build_docx(self._DOCS, title="공문서 제목", gov_options=gov_options)
        doc = Document(io.BytesIO(raw))
        para_text = "\n".join(p.text for p in doc.paragraphs)
        # Approval block is in a table — also scan table cells
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        return para_text + "\n" + table_text

    def test_gov_header_contains_doc_number(self):
        opts = GovDocOptions(
            doc_number="행안부-2025-001",
            is_government_format=True,
        )
        text = self._full_text(gov_options=opts)
        assert "행안부-2025-001" in text

    def test_gov_header_contains_recipient(self):
        opts = GovDocOptions(recipient="수신자 참조", is_government_format=True)
        text = self._full_text(gov_options=opts)
        assert "수신자 참조" in text

    def test_gov_header_contains_title(self):
        opts = GovDocOptions(is_government_format=True)
        text = self._full_text(gov_options=opts)
        assert "공문서 제목" in text

    def test_gov_header_contains_org_name(self):
        opts = GovDocOptions(org_name="행정안전부", is_government_format=True)
        text = self._full_text(gov_options=opts)
        assert "행정안전부" in text

    def test_approval_block_shows_approver(self):
        opts = GovDocOptions(
            approver="이영희 국장",
            drafter="홍길동",
            is_government_format=True,
        )
        text = self._full_text(gov_options=opts)
        assert "이영희 국장" in text
        assert "홍길동" in text

    def test_attachments_shown(self):
        opts = GovDocOptions(
            attachments=["계획서 1부", "예산안 1부"],
            is_government_format=True,
        )
        text = self._full_text(gov_options=opts)
        assert "계획서 1부" in text

    def test_no_gov_block_when_flag_false(self):
        opts = GovDocOptions(doc_number="행안부-999", is_government_format=False)
        text = self._full_text(gov_options=opts)
        # doc_number should NOT appear when flag is False
        assert "행안부-999" not in text


# ===========================================================================
# D. HWP — style definitions
# ===========================================================================

class TestHwpStyles:
    def test_styles_xml_contains_본문(self):
        xml = _styles_xml()
        assert "본문" in xml

    def test_styles_xml_contains_제목1(self):
        xml = _styles_xml()
        assert "제목1" in xml

    def test_styles_xml_contains_제목2(self):
        xml = _styles_xml()
        assert "제목2" in xml

    def test_styles_xml_contains_제목3(self):
        xml = _styles_xml()
        assert "제목3" in xml

    def test_styles_xml_not_empty(self):
        """Old bug: <hh:styles/> was empty. Now must have <hh:style> elements."""
        xml = _styles_xml()
        assert "<hh:style " in xml

    def test_styles_xml_font_name(self):
        xml = _styles_xml(font_name="나눔고딕")
        assert "나눔고딕" in xml

    def test_styles_xml_font_size_1050(self):
        """10.5pt should produce hangul=1050."""
        xml = _styles_xml(font_size_pt=10.5)
        assert 'hh:hangul="1050"' in xml

    def test_header_xml_styles_not_empty_tag(self):
        """header.xml must NOT contain empty <hh:styles/> self-closing tag."""
        raw = build_hwp([{"doc_type": "t", "markdown": "# 제목"}], title="X")
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            header = zf.read("Contents/header.xml").decode()
        assert "<hh:styles/>" not in header
        assert "<hh:styles " in header or "<hh:styles>" in header


# ===========================================================================
# E. HWP — page layout (secPr)
# ===========================================================================

class TestHwpPageLayout:
    _DOCS = [{"doc_type": "t", "markdown": "내용"}]

    def _section(self, **kwargs):
        return _section_xml(self._DOCS, "제목", None, **kwargs)

    def test_secpr_present(self):
        xml = _section_xml(self._DOCS, "제목", None, 30, 15, 20, 20)
        assert "<hh:secPr>" in xml

    def test_a4_width_in_secpr(self):
        xml = _section_xml(self._DOCS, "제목", None, 30, 15, 20, 20)
        assert f'hh:width="{_A4_W}"' in xml

    def test_a4_height_in_secpr(self):
        xml = _section_xml(self._DOCS, "제목", None, 30, 15, 20, 20)
        assert f'hh:height="{_A4_H}"' in xml

    def test_custom_margin_in_secpr(self):
        from app.services.hwp_service import _mm
        xml = _section_xml(self._DOCS, "제목", None, 40, 25, 25, 25)
        top_val = _mm(40)
        assert f'hh:top="{top_val}"' in xml

    def test_hwp_zip_valid(self):
        raw = build_hwp(self._DOCS, title="테스트")
        assert zipfile.is_zipfile(io.BytesIO(raw))

    def test_hwp_has_section0(self):
        raw = build_hwp(self._DOCS, title="테스트")
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            assert "Contents/section0.xml" in zf.namelist()


# ===========================================================================
# F. HWP — gov format content
# ===========================================================================

class TestHwpGovFormat:
    _DOCS = [{"doc_type": "t", "markdown": "본문 내용"}]

    def _section_content(self, opts=None):
        from app.services.hwp_service import _section_xml
        top = opts.top_margin_mm if opts else 30
        bot = opts.bottom_margin_mm if opts else 15
        left = opts.left_margin_mm if opts else 20
        right = opts.right_margin_mm if opts else 20
        return _section_xml(self._DOCS, "공문서 제목", opts, top, bot, left, right)

    def test_gov_header_doc_number(self):
        opts = GovDocOptions(doc_number="XXXX-001", is_government_format=True)
        xml = self._section_content(opts)
        assert "XXXX-001" in xml

    def test_gov_header_org_name(self):
        opts = GovDocOptions(org_name="행정안전부", is_government_format=True)
        xml = self._section_content(opts)
        assert "행정안전부" in xml

    def test_approval_block_text(self):
        opts = GovDocOptions(approver="김과장", is_government_format=True)
        xml = self._section_content(opts)
        assert "김과장" in xml

    def test_no_gov_block_when_flag_false(self):
        opts = GovDocOptions(doc_number="YYYY-999", is_government_format=False)
        xml = self._section_content(opts)
        assert "YYYY-999" not in xml


# ===========================================================================
# G. PDF — CSS margin generation
# ===========================================================================

class TestPdfCss:
    def test_default_margin_in_css(self):
        css = _build_css(None)
        assert "20mm" in css

    def test_custom_top_margin(self):
        opts = GovDocOptions(top_margin_mm=35)
        css = _build_css(opts)
        assert "35mm" in css

    def test_custom_font_in_css(self):
        opts = GovDocOptions(font_name="나눔고딕")
        css = _build_css(opts)
        assert "나눔고딕" in css

    def test_font_size_in_css(self):
        opts = GovDocOptions(font_size_pt=12.0)
        css = _build_css(opts)
        assert "12.0pt" in css


# ===========================================================================
# H. PDF — gov header block / approval block HTML
# ===========================================================================

class TestPdfGovHtml:
    def test_gov_header_html_contains_doc_number(self):
        opts = GovDocOptions(doc_number="행안부-9999", is_government_format=True)
        html = _gov_header_block_html("제목", opts)
        assert "행안부-9999" in html

    def test_gov_header_html_contains_recipient(self):
        opts = GovDocOptions(recipient="홍길동 부장", is_government_format=True)
        html = _gov_header_block_html("제목", opts)
        assert "홍길동 부장" in html

    def test_gov_header_html_contains_org(self):
        opts = GovDocOptions(org_name="국토교통부", is_government_format=True)
        html = _gov_header_block_html("제목", opts)
        assert "국토교통부" in html

    def test_approval_block_html_shows_roles(self):
        opts = GovDocOptions(drafter="기안자A", reviewer="검토자B", approver="결재자C")
        html = _approval_block_html(opts)
        assert "기안자A" in html
        assert "검토자B" in html
        assert "결재자C" in html

    def test_approval_block_html_empty_when_no_approvers(self):
        opts = GovDocOptions()  # no approvers set
        html = _approval_block_html(opts)
        assert html == ""

    def test_full_render_includes_gov_block(self):
        opts = GovDocOptions(doc_number="TEST-001", is_government_format=True)
        docs = [{"doc_type": "t", "markdown": "내용"}]
        html = _render_html(docs, "공문 제목", opts)
        assert "TEST-001" in html
        assert "공문 제목" in html


# ===========================================================================
# I. API — EditedExportRequest schema accepts gov_options
# ===========================================================================

class TestEditedExportRequestSchema:
    def test_gov_options_none_by_default(self):
        req = EditedExportRequest(
            format="docx",
            docs=[EditedDocInput(doc_type="t", markdown="내용")],
        )
        assert req.gov_options is None

    def test_gov_options_accepts_dict(self):
        req = EditedExportRequest(
            format="hwp",
            docs=[EditedDocInput(doc_type="t", markdown="내용")],
            gov_options={"is_government_format": True, "org_name": "행안부"},
        )
        assert req.gov_options["org_name"] == "행안부"

    def test_gov_options_not_required(self):
        """Omitting gov_options must not raise a validation error."""
        data = {"format": "docx", "docs": [{"doc_type": "t", "markdown": "x"}]}
        req = EditedExportRequest(**data)
        assert req.gov_options is None


# ===========================================================================
# J. API endpoint — export-edited with gov_options (via test client)
# ===========================================================================

@pytest.fixture
def client():
    from fastapi.testclient import TestClient
    from app.main import create_app
    return TestClient(create_app())


class TestExportEditedGovOptions:
    _BODY_BASE = {
        "format": "docx",
        "title": "공문서 테스트",
        "docs": [{"doc_type": "t", "markdown": "# 제목\n\n내용입니다."}],
    }

    def test_export_edited_docx_no_gov_options(self, client):
        res = client.post("/generate/export-edited", json=self._BODY_BASE)
        assert res.status_code == 200
        assert res.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_export_edited_docx_with_gov_options(self, client):
        body = {**self._BODY_BASE, "gov_options": {
            "is_government_format": True,
            "org_name": "행정안전부",
            "doc_number": "행안부-2025-TEST",
            "drafter": "홍길동",
            "approver": "김철수",
        }}
        res = client.post("/generate/export-edited", json=body)
        assert res.status_code == 200
        # Verify the DOCX content includes the gov doc header
        raw = res.content
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "행정안전부" in text
        assert "행안부-2025-TEST" in text

    def test_export_edited_hwp_with_gov_options(self, client):
        body = {
            "format": "hwp",
            "title": "HWP 공문서",
            "docs": [{"doc_type": "t", "markdown": "내용"}],
            "gov_options": {
                "is_government_format": True,
                "org_name": "국토교통부",
            },
        }
        res = client.post("/generate/export-edited", json=body)
        assert res.status_code == 200
        assert res.headers["content-type"] == "application/hwp+zip"
        raw = res.content
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            section_xml = zf.read("Contents/section0.xml").decode()
        assert "국토교통부" in section_xml

    def test_export_edited_invalid_gov_options_graceful(self, client):
        """Invalid gov_options fields should be ignored gracefully (not 500)."""
        body = {**self._BODY_BASE, "gov_options": {"unknown_field": "x"}}
        res = client.post("/generate/export-edited", json=body)
        # Should return 200 (unknown fields cause gov_options to be None via _resolve_gov_options)
        assert res.status_code == 200


# ===========================================================================
# K. _resolve_gov_options helper
# ===========================================================================

class TestResolveGovOptions:
    def test_none_input_returns_none(self):
        from app.routers.generate import _resolve_gov_options  # type: ignore[attr-defined]
        assert _resolve_gov_options(None) is None

    def test_empty_dict_returns_none(self):
        from app.routers.generate import _resolve_gov_options  # type: ignore[attr-defined]
        assert _resolve_gov_options({}) is None

    def test_valid_dict_returns_gov_doc_options(self):
        from app.routers.generate import _resolve_gov_options  # type: ignore[attr-defined]
        opts = _resolve_gov_options({"org_name": "행안부", "is_government_format": True})
        assert opts is not None
        assert opts.org_name == "행안부"
        assert opts.is_government_format is True

    def test_invalid_fields_returns_none(self):
        from app.routers.generate import _resolve_gov_options  # type: ignore[attr-defined]
        opts = _resolve_gov_options({"nonexistent_field": "value"})
        assert opts is None
