"""Tests for POST /generate/docx endpoint and build_docx() service."""
from __future__ import annotations

import zipfile
from io import BytesIO

from fastapi.testclient import TestClient

_DOCX_MAGIC = b"PK\x03\x04"  # OOXML/ZIP magic bytes — all .docx files start with this


def _create_client(tmp_path, monkeypatch):
    monkeypatch.setenv("DECISIONDOC_PROVIDER", "mock")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("DECISIONDOC_ENV", "dev")
    monkeypatch.setenv("DECISIONDOC_MAINTENANCE", "0")
    monkeypatch.delenv("DECISIONDOC_API_KEY", raising=False)
    monkeypatch.delenv("DECISIONDOC_API_KEYS", raising=False)
    from app.main import create_app

    return TestClient(create_app())


# ── Unit tests for build_docx() ──────────────────────────────────────────────

def test_build_docx_returns_valid_ooxml_bytes():
    """build_docx() must return bytes that start with the ZIP/OOXML magic."""
    from app.services.docx_service import build_docx

    result = build_docx(
        [{"doc_type": "adr", "markdown": "# 제목\n\n본문 내용"}],
        title="테스트 문서",
    )
    assert result[:4] == _DOCX_MAGIC


def test_build_docx_handles_bold_inline():
    """**bold** spans must not crash the builder."""
    from app.services.docx_service import build_docx

    result = build_docx(
        [{"doc_type": "x", "markdown": "**굵은 텍스트** 일반 텍스트"}],
        title="볼드 테스트",
    )
    assert result[:4] == _DOCX_MAGIC


def test_build_docx_multiple_docs():
    """Multiple docs must produce a non-empty bytes result (page breaks included)."""
    from app.services.docx_service import build_docx

    docs = [
        {"doc_type": "adr",      "markdown": "# 결정\n\n내용 A"},
        {"doc_type": "onepager", "markdown": "## 요약\n\n- 항목 1\n- 항목 2"},
    ]
    result = build_docx(docs, title="멀티 문서")
    assert result[:4] == _DOCX_MAGIC
    assert len(result) > 1000  # must be a real docx, not empty


def test_build_docx_renders_markdown_tables():
    """Markdown tables must become real Word tables in the exported document."""
    from docx import Document

    from app.services.docx_service import build_docx

    markdown = (
        "# 문서 제목\n\n"
        "| 단계 | 기간 | 주요 산출물 | 마일스톤 |\n"
        "| --- | --- | --- | --- |\n"
        "| 착수 | 1개월 | 착수보고서 | M1 |\n"
        "| 개발 | 4개월 | 개발 산출물 | M2 |\n"
    )
    result = build_docx([{"doc_type": "x", "markdown": markdown}], title="표 테스트")
    doc = Document(BytesIO(result))

    assert len(doc.tables) >= 1
    target = next(
        table for table in doc.tables
        if table.cell(0, 0).text == "단계" and table.cell(0, 1).text == "기간"
    )
    assert target.cell(1, 0).text == "착수"


def test_build_docx_adds_export_cover_and_section_intro():
    """Non-government DOCX exports should include a cover page and per-doc section intro."""
    from docx import Document

    from app.services.docx_service import build_docx

    result = build_docx(
        [
            {"doc_type": "business_understanding", "markdown": "# 제목\n\n본문 A"},
            {"doc_type": "tech_proposal", "markdown": "# 제목\n\n본문 B"},
        ],
        title="완성형 패키지 테스트",
    )
    doc = Document(BytesIO(result))
    paragraph_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    )
    joined = "\n".join([paragraph_text, table_text])

    assert "완성형 문서 패키지" in joined
    assert "문서 구성" in joined
    assert "핵심 검토 포인트" in joined
    assert "문서 수" in joined
    assert "표 수" in joined
    assert "사업 이해" in joined
    assert "문서 01 / 02" in joined
    assert "검토 초점" in joined
    assert "핵심 섹션:" in joined
    assert "구성 지표:" in joined
    assert " / 구성 특징:" not in joined


def test_build_docx_embeds_generated_visual_asset_preview():
    from docx import Document

    from app.services.docx_service import build_docx

    docs = [
        {
            "doc_type": "business_understanding",
            "markdown": "# 제목\n\n본문 A",
        }
    ]
    visual_assets = [
        {
            "asset_id": "asset-1",
            "doc_type": "business_understanding",
            "slide_title": "사업 추진 배경",
            "visual_type": "현장 사진",
            "visual_brief": "운영 현장을 보여주는 생성 이미지",
            "media_type": "image/png",
            "content_base64": "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO5tm8sAAAAASUVORK5CYII=",
        }
    ]

    result = build_docx(docs, title="시각자료 테스트", visual_assets=visual_assets)
    doc = Document(BytesIO(result))
    joined = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    assert "생성 시각자료" in joined
    assert "사업 추진 배경" in joined
    with zipfile.ZipFile(BytesIO(result)) as zf:
        media_files = [name for name in zf.namelist() if name.startswith("word/media/")]
    assert media_files


# ── Integration tests for /generate/docx ─────────────────────────────────────

def test_docx_endpoint_returns_binary(tmp_path, monkeypatch):
    """Endpoint must return valid DOCX binary with 200 status."""
    client = _create_client(tmp_path, monkeypatch)
    res = client.post("/generate/docx", json={"title": "Word 테스트", "goal": "검증"})
    assert res.status_code == 200
    assert res.content[:4] == _DOCX_MAGIC


def test_docx_content_type(tmp_path, monkeypatch):
    """Response Content-Type must indicate Word document format."""
    client = _create_client(tmp_path, monkeypatch)
    res = client.post("/generate/docx", json={"title": "t", "goal": "g"})
    assert "wordprocessingml" in res.headers["content-type"]


def test_docx_content_disposition_has_attachment(tmp_path, monkeypatch):
    """Content-Disposition must include 'attachment' and RFC 5987 filename."""
    client = _create_client(tmp_path, monkeypatch)
    res = client.post("/generate/docx", json={"title": "한글 제목", "goal": "g"})
    cd = res.headers.get("content-disposition", "")
    assert "attachment" in cd
    assert "filename*=UTF-8" in cd


def test_docx_all_bundles_supported(tmp_path, monkeypatch):
    """Every bundle type must return a valid DOCX (no BUNDLE_NOT_SUPPORTED)."""
    client = _create_client(tmp_path, monkeypatch)
    for bundle_type in ["tech_decision", "proposal_kr", "presentation_kr"]:
        res = client.post(
            "/generate/docx",
            json={"title": "t", "goal": "g", "bundle_type": bundle_type},
        )
        assert res.status_code == 200, f"Failed for bundle_type={bundle_type}"
        assert res.content[:4] == _DOCX_MAGIC
